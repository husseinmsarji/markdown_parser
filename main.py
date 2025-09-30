from __future__ import annotations

import base64
import os
import re
from io import BytesIO
from typing import Dict, List, Optional, Tuple

# --- third-party (see notes above) ---
try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.shared import Inches, Pt
except ImportError as e:
    raise ImportError(
        "Missing dependency 'python-docx'. Install via: pip install python-docx"
    ) from e

try:
    from markdown_it import MarkdownIt
    from mdit_py_plugins.footnote import footnote_plugin
    from mdit_py_plugins.tasklists import tasklists_plugin
    from mdit_py_plugins.table import table_plugin
    from mdit_py_plugins.deflist import deflist_plugin
except ImportError as e:
    raise ImportError(
        "Missing dependency 'markdown-it-py' (and plugins). Install via: "
        "pip install markdown-it-py mdit-py-plugins"
    ) from e

try:
    import requests  # optional
except Exception:  # pragma: no cover
    requests = None


# --------------------------- public API ---------------------------

def markdown_to_docx(markdown_text: str, output_path: str) -> str:
    """
    Convert an LLM-generated Markdown string to a cleaned, readable .docx file.

    Parameters
    ----------
    markdown_text : str
        The raw Markdown string (possibly wrapped in ```markdown fences or with YAML front-matter).
    output_path : str
        The desired .docx output path. The function creates parent directories as needed.
        If output_path lacks the .docx suffix, it will be appended.

    Returns
    -------
    str
        The absolute path to the written .docx file.

    Notes
    -----
    - Robustly strips common LLM wrappers (e.g., a top-level ```markdown ... ``` block).
    - Supports: headings, paragraphs, emphasis/strong/strike, inline code, code fences,
      blockquotes, bullet/numbered lists (with nesting), task lists, definition lists,
      tables, links (rendered as real docx hyperlinks), images (embedded when possible),
      horizontal rules, and footnotes (rendered as an end section titled "Footnotes").
    - If 'requests' is unavailable or an image fetch fails, the image is represented by
      its alt text and URL.
    """
    if not isinstance(markdown_text, str):
        raise TypeError("markdown_text must be a string")
    if not isinstance(output_path, str):
        raise TypeError("output_path must be a string")

    text = _preprocess_markdown(markdown_text)

    # --- parse markdown into tokens ---
    md = (
        MarkdownIt("commonmark", options_update={"linkify": True, "typographer": True})
        .use(table_plugin)
        .use(tasklists_plugin, label=True)
        .use(deflist_plugin)
        .use(footnote_plugin)
    )
    env: Dict = {}
    tokens = md.parse(text, env)

    # --- build the Word document ---
    doc = Document()
    _ensure_minimum_styles(doc)

    list_stack: List[str] = []  # sequence of 'ul' or 'ol' to track nesting
    in_blockquote = 0
    footnote_defs: Dict[int, str] = {}
    seen_footnote_ids: List[int] = []

    i = 0
    n = len(tokens)

    while i < n:
        t = tokens[i]

        # Footnote definitions are emitted in a special block; parse & stash
        if t.type == "footnote_block_open":
            i = _consume_footnote_block(tokens, i, footnote_defs, md)
            continue

        # Headings
        if t.type == "heading_open":
            level = int(t.tag[1])
            inline = tokens[i + 1] if i + 1 < n and tokens[i + 1].type == "inline" else None
            if inline:
                p = doc.add_heading("", level=level)
                _render_inlines(doc, p, inline.children or [], seen_footnote_ids)
            i += 3  # open -> inline -> close
            continue

        # Paragraphs (top-level or inside lists / blockquotes)
        if t.type == "paragraph_open":
            style = "Intense Quote" if in_blockquote else None
            inline = tokens[i + 1] if i + 1 < n and tokens[i + 1].type == "inline" else None
            if inline:
                p = doc.add_paragraph(style=style)
                _render_inlines(doc, p, inline.children or [], seen_footnote_ids)
            i += 3
            continue

        # Lists
        if t.type == "bullet_list_open":
            list_stack.append("ul")
            i += 1
            continue

        if t.type == "ordered_list_open":
            list_stack.append("ol")
            i += 1
            continue

        if t.type == "bullet_list_close" or t.type == "ordered_list_close":
            if list_stack:
                list_stack.pop()
            i += 1
            continue

        if t.type == "list_item_open":
            # Find first inline chunk within this list item for the main bullet text
            j = i + 1
            end = _find_matching(tokens, j, "list_item_close")
            bullet_inline = None
            # support items that are simple text or a paragraph
            k = j
            while k < end:
                if tokens[k].type == "paragraph_open" and tokens[k + 1].type == "inline":
                    bullet_inline = tokens[k + 1]
                    break
                if tokens[k].type == "inline":
                    bullet_inline = tokens[k]
                    break
                k += 1

            style = "List Bullet" if (list_stack and list_stack[-1] == "ul") else "List Number"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Inches(0.25 * max(0, len(list_stack) - 1))

            # Task list marker support ([ ] / [x]) even if plugin didn't tag
            if bullet_inline and bullet_inline.children:
                _maybe_render_task_marker_prefix(p, bullet_inline)

            if bullet_inline:
                _render_inlines(doc, p, bullet_inline.children or [], seen_footnote_ids)

            # Process any nested blocks inside the list item (e.g., sublists)
            k = j
            while k < end:
                tk = tokens[k]
                if tk.type in ("bullet_list_open", "ordered_list_open"):
                    # Nested lists are handled by the main loop once we advance index
                    pass
                elif tk.type == "paragraph_open":
                    # Already handled by main bullet text; skip the body paragraph to avoid duplication
                    pass
                k += 1

            i = end + 1
            continue

        # Blockquotes
        if t.type == "blockquote_open":
            in_blockquote += 1
            i += 1
            continue
        if t.type == "blockquote_close":
            in_blockquote = max(0, in_blockquote - 1)
            i += 1
            continue

        # Code blocks (fences and indented)
        if t.type in ("fence", "code_block"):
            lang = (t.info or "").strip() if t.type == "fence" else ""
            code = (t.content or "").rstrip("\n")
            if lang:
                doc.add_paragraph(f"{lang}:", style="Intense Emphasis")
            _add_code_block_paragraph(doc, code)
            i += 1
            continue

        # Horizontal rule
        if t.type == "hr":
            doc.add_paragraph("—" * 24)
            i += 1
            continue

        # Tables
        if t.type == "table_open":
            i = _consume_table(tokens, i, doc)
            continue

        # Definition lists (term/definition)
        if t.type == "dl_open":
            i = _consume_deflist(tokens, i, doc, md)
            continue

        # Anything else: step forward
        i += 1

    # If any footnotes were referenced, append them as an end section
    if seen_footnote_ids and footnote_defs:
        doc.add_page_break()
        doc.add_heading("Footnotes", level=1)
        for fid in sorted(dict.fromkeys(seen_footnote_ids)):
            txt = footnote_defs.get(fid, "")
            p = doc.add_paragraph()
            p.add_run(f"[{fid + 1}] ").bold = True
            p.add_run(txt)

    # --- write the file ---
    output_path = _normalize_output_path(output_path)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return os.path.abspath(output_path)


# --------------------------- helpers ---------------------------

def _preprocess_markdown(text: str) -> str:
    """Normalize newlines, strip BOM, remove YAML front matter and unwrap a single outer ```markdown fence."""
    if not text:
        return ""

    # Normalize newlines and trim BOM/whitespace
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.strip("\ufeff").strip()

    # Strip top YAML front matter if present
    if text.startswith("---\n"):
        m = re.match(r"^---\n.*?\n---\n(.*)$", text, flags=re.S)
        if m:
            text = m.group(1).lstrip()

    # Unwrap outer triple backticks if they declare markdown (or empty but clearly markdowny)
    m = re.match(r"^\s*```([^\n]*)\n(.*)\n```\s*$", text, flags=re.S)
    if m:
        lang = (m.group(1) or "").strip().lower()
        inner = m.group(2)
        if lang in ("markdown", "md"):
            text = inner.strip()
        elif lang == "" and _looks_like_markdown(inner):
            text = inner.strip()

    # Drop stray HTML tags while preserving text content (very light sanitize)
    # This avoids raw <details>, <summary>, etc. bleeding into the doc.
    text = re.sub(r"<\s*br\s*/?\s*>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)

    return text


def _looks_like_markdown(s: str) -> bool:
    """Heuristic for deciding if an outer fence likely wraps markdown, not code."""
    hits = 0
    patterns = [
        r"^#{1,6}\s",     # headings
        r"^\s*[-*]\s+",   # bullets
        r"^\s*\d+\.\s+",  # numbered list
        r"\|.+\|",        # table row
        r"\[.+?\]\(.+?\)",# link
        r"```",           # inner code blocks
        r"^>\s",          # blockquote
    ]
    for pat in patterns:
        if re.search(pat, s, flags=re.M):
            hits += 1
    return hits >= 2


def _ensure_minimum_styles(doc: Document) -> None:
    """Create a couple of helper styles if they don't exist."""
    styles = doc.styles

    if "CodeBlock" not in styles:
        style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(10)

    if "InlineCode" not in styles:
        style = styles.add_style("InlineCode", WD_STYLE_TYPE.CHARACTER)
        style.font.name = "Consolas"
        style.font.size = Pt(10)


def _add_code_block_paragraph(doc: Document, code: str) -> None:
    p = doc.add_paragraph(style="CodeBlock")
    # Preserve line breaks
    lines = code.split("\n")
    for idx, line in enumerate(lines):
        r = p.add_run(line)
        if idx < len(lines) - 1:
            r.add_break()


def _render_inlines(doc: Document, paragraph, children, seen_footnote_ids: List[int]) -> None:
    """
    Render inline tokens into a given paragraph.
    Handles text/emphasis/strong/strike/inline code/links/images/footnote refs/line breaks.
    """
    # current text state toggles
    state = {"bold": False, "italic": False, "strike": False, "underline": False}

    i = 0
    while i < len(children):
        tok = children[i]

        if tok.type == "text":
            _add_text_run(paragraph, tok.content or "", state)

        elif tok.type == "softbreak":
            paragraph.add_run().add_break()

        elif tok.type == "hardbreak":
            paragraph.add_run().add_break()

        elif tok.type == "code_inline":
            r = paragraph.add_run(tok.content or "")
            r.style = paragraph.part.document.styles["InlineCode"]

        elif tok.type == "em_open":
            state["italic"] = True
        elif tok.type == "em_close":
            state["italic"] = False

        elif tok.type == "strong_open":
            state["bold"] = True
        elif tok.type == "strong_close":
            state["bold"] = False

        elif tok.type in ("s_open", "strikethrough_open"):
            state["strike"] = True
        elif tok.type in ("s_close", "strikethrough_close"):
            state["strike"] = False

        elif tok.type == "link_open":
            href = tok.attrGet("href") or ""
            # gather inner plain text
            j = i + 1
            inner_text = []
            while j < len(children) and children[j].type != "link_close":
                cj = children[j]
                if cj.type in ("text", "code_inline"):
                    inner_text.append(cj.content)
                elif cj.type == "softbreak":
                    inner_text.append("\n")
                elif cj.type == "image":
                    # use alt text inside link
                    inner_text.append(cj.attrGet("alt") or "")
                j += 1
            link_text = "".join(inner_text).strip() or href
            _add_hyperlink(paragraph, href, link_text, state)
            i = j  # skip to link_close; the loop will add +1
        elif tok.type == "image":
            src = tok.attrGet("src") or ""
            alt = tok.attrGet("alt") or ""
            _add_image_block(doc, src, alt, paragraph_after=True)

        elif tok.type == "footnote_ref":
            fid = tok.meta["id"] if tok.meta and "id" in tok.meta else tok.content
            try:
                fid_int = int(fid)
            except Exception:
                fid_int = int(fid) if isinstance(fid, int) else 0
            seen_footnote_ids.append(fid_int)
            r = paragraph.add_run(f"[{fid_int + 1}]")
            r.font.superscript = True

        # ignore other inline types safely
        i += 1


def _add_text_run(paragraph, text: str, state: Dict[str, bool]) -> None:
    if not text:
        return
    r = paragraph.add_run(text)
    r.bold = state.get("bold", False)
    r.italic = state.get("italic", False)
    r.font.strike = state.get("strike", False)
    r.underline = state.get("underline", False)


def _add_hyperlink(paragraph, url: str, text: str, state: Dict[str, bool]) -> None:
    """Insert a real docx hyperlink (blue & underlined), preserving bold/italic/strike."""
    # guard against empty/invalid URLs
    url = (url or "").strip()
    if not url:
        _add_text_run(paragraph, text, state)
        return

    # relationship id
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # <w:hyperlink r:id="..."><w:r><w:rPr>...</w:rPr><w:t>text</w:t></w:r></w:hyperlink>
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    if state.get("bold"):
        rPr.append(OxmlElement("w:b"))
    if state.get("italic"):
        rPr.append(OxmlElement("w:i"))
    if state.get("strike"):
        rPr.append(OxmlElement("w:strike"))

    t = OxmlElement("w:t")
    # Ensure xml preserves whitespace if needed
    t.set(qn("xml:space"), "preserve")
    t.text = text

    new_run.append(rPr)
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_image_block(doc: Document, src: str, alt: str, paragraph_after: bool = True) -> None:
    """
    Try to embed image from src. Supports http(s), data URIs, and local filesystem paths.
    If embedding fails, writes 'alt (src)' as a paragraph instead.
    """
    img_bytes: Optional[BytesIO] = None
    try:
        if src.startswith("data:image/"):
            # data URI
            comma = src.find(",")
            if comma != -1:
                b64 = src[comma + 1 :]
                img_bytes = BytesIO(base64.b64decode(b64))
        elif src.startswith("http://") or src.startswith("https://"):
            if requests is not None:
                resp = requests.get(src, timeout=8)
                resp.raise_for_status()
                img_bytes = BytesIO(resp.content)
        else:
            # local file path
            if os.path.exists(src) and os.path.isfile(src):
                with open(src, "rb") as f:
                    img_bytes = BytesIO(f.read())
    except Exception:
        img_bytes = None

    if img_bytes:
        pic = doc.add_picture(img_bytes)  # natural size; or use width=Inches(6)
        # (Optional) add a caption-like line using alt text
        if alt:
            p = doc.add_paragraph()
            p.add_run(alt).italic = True
        if paragraph_after:
            doc.add_paragraph()
    else:
        # graceful fallback: text with URL
        p = doc.add_paragraph()
        fallback = alt.strip() or "Image"
        p.add_run(fallback + " ").italic = True
        if src:
            _add_hyperlink(p, src, f"({src})", {"bold": False, "italic": False, "strike": False, "underline": True})


def _maybe_render_task_marker_prefix(paragraph, inline_token) -> None:
    """
    Detect a task-list marker at the very start of the item's text and render a checkbox prefix.
    E.g., "[ ] item" or "[x] item" -> "☐" / "☑" prefix.
    Also trims the marker from the first text node so it isn't duplicated.
    """
    children = inline_token.children or []
    if not children or children[0].type != "text":
        return
    txt = children[0].content or ""
    m = re.match(r"^\s*\[([ xX])\]\s+", txt)
    if m:
        checked = m.group(1).lower() == "x"
        paragraph.add_run("☑ " if checked else "☐ ")
        # trim the marker from the first text node
        children[0].content = txt[m.end() :]


def _find_matching(tokens, start_idx: int, closing_type: str) -> int:
    """Find the index of the next token with type == closing_type (same nesting level)."""
    for k in range(start_idx, len(tokens)):
        if tokens[k].type == closing_type:
            return k
    return len(tokens) - 1


def _consume_table(tokens, i: int, doc: Document) -> int:
    """
    Parse a markdown-it table starting at tokens[i] == 'table_open' and append a docx table.
    Returns the index right after 'table_close'.
    """
    # Extract rows (header + body) as plain text
    headers: List[str] = []
    rows: List[List[str]] = []

    k = i + 1
    while k < len(tokens) and tokens[k].type != "table_close":
        t = tokens[k]

        if t.type == "thead_open":
            # one header row (tr)
            k += 1
            while tokens[k].type != "thead_close":
                if tokens[k].type == "tr_open":
                    k += 1
                    row: List[str] = []
                    while tokens[k].type != "tr_close":
                        if tokens[k].type == "th_open":
                            # th_open -> inline -> th_close
                            inline = tokens[k + 1]
                            row.append(_inline_text(inline.children or []))
                            k += 3
                        else:
                            k += 1
                    headers = row
                    k += 1
                else:
                    k += 1
            k += 1
            continue

        if t.type == "tbody_open":
            k += 1
            while tokens[k].type != "tbody_close":
                if tokens[k].type == "tr_open":
                    k += 1
                    row: List[str] = []
                    while tokens[k].type != "tr_close":
                        if tokens[k].type == "td_open":
                            inline = tokens[k + 1]
                            row.append(_inline_text(inline.children or []))
                            k += 3
                        else:
                            k += 1
                    rows.append(row)
                    k += 1
                else:
                    k += 1
            k += 1
            continue

        k += 1

    # Build the docx table
    cols = max(len(headers), max((len(r) for r in rows), default=0))
    cols = max(cols, 1)
    table = doc.add_table(rows=(1 if headers else 0) + len(rows), cols=cols)
    table.style = "Table Grid"

    r = 0
    if headers:
        for c, h in enumerate(headers):
            cell = table.cell(0, c)
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
        r = 1

    for row in rows:
        for c in range(cols):
            cell = table.cell(r, c)
            cell.text = row[c] if c < len(row) else ""
        r += 1

    # return index after 'table_close'
    while i < len(tokens) and tokens[i].type != "table_close":
        i += 1
    return i + 1


def _consume_deflist(tokens, i: int, doc: Document, md: MarkdownIt) -> int:
    """
    Parse a markdown-it definition list starting at 'dl_open' and append as paragraphs:
    Term: definition...
    Returns the index right after 'dl_close'.
    """
    k = i + 1
    current_term: Optional[str] = None
    while k < len(tokens) and tokens[k].type != "dl_close":
        t = tokens[k]
        if t.type == "dt_open":
            inline = tokens[k + 1]  # inline
            current_term = _inline_text(inline.children or [])
            k += 3  # dt_open, inline, dt_close
            continue
        if t.type == "dd_open":
            # the dd can contain multiple paragraphs/blocks before dd_close
            k += 1
            parts: List[str] = []
            while tokens[k].type != "dd_close":
                if tokens[k].type == "paragraph_open" and tokens[k + 1].type == "inline":
                    parts.append(_inline_text(tokens[k + 1].children or []))
                    k += 3
                else:
                    # Fallback: try to serialize anything else
                    if tokens[k].type == "inline":
                        parts.append(_inline_text(tokens[k].children or []))
                    k += 1
            definition = "\n".join(p for p in parts if p)
            p = doc.add_paragraph()
            if current_term:
                p.add_run(f"{current_term}: ").bold = True
            p.add_run(definition)
            k += 1  # past dd_close
            continue
        k += 1

    # return index after 'dl_close'
    while i < len(tokens) and tokens[i].type != "dl_close":
        i += 1
    return i + 1


def _inline_text(children) -> str:
    """Flatten inline tokens into plain text."""
    out: List[str] = []
    for t in children:
        if t.type in ("text", "code_inline"):
            out.append(t.content or "")
        elif t.type == "softbreak":
            out.append("\n")
        elif t.type == "image":
            out.append(t.attrGet("alt") or "")
    return "".join(out)


def _consume_footnote_block(tokens, i: int, footnote_defs: Dict[int, str], md: MarkdownIt) -> int:
    """
    Collect footnote definitions emitted by mdit footnote plugin.
    They appear as:
        footnote_block_open
          footnote_open (meta.id=k)
            paragraph_open
              inline ...
            paragraph_close
          footnote_close
        footnote_block_close
    """
    k = i + 1
    current_id: Optional[int] = None
    buffer: List[str] = []
    while k < len(tokens) and tokens[k].type != "footnote_block_close":
        t = tokens[k]
        if t.type == "footnote_open":
            current_id = int(t.meta["id"])
            buffer = []
            k += 1
            continue
        if t.type == "footnote_close":
            if current_id is not None:
                footnote_defs[current_id] = "\n".join(buffer).strip()
            current_id = None
            buffer = []
            k += 1
            continue
        if t.type == "paragraph_open" and tokens[k + 1].type == "inline":
            buffer.append(_inline_text(tokens[k + 1].children or []))
            k += 3
            continue
        # Other blocks inside footnotes (rare): just skip or flatten
        if t.type == "inline":
            buffer.append(_inline_text(t.children or []))
        k += 1

    # return index after block close
    while i < len(tokens) and tokens[i].type != "footnote_block_close":
        i += 1
    return i + 1


def _normalize_output_path(path: str) -> str:
    path = os.path.expanduser(path)
    if not path.lower().endswith(".docx"):
        path = path + ".docx"
    if os.path.isdir(path):
        raise ValueError("output_path points to a directory; please provide a file path")
    parent = os.path.dirname(path) or "."
    os.makedirs(parent, exist_ok=True)
    return path


# --------------------------- end helpers ---------------------------
